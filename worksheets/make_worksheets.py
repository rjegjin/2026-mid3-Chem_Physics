#!/usr/bin/env python3
"""15·16강 학습지를 Google Docs로 열 수 있는 .docx로 만든다.

WORKSHEET_CREATION_WORKFLOW.md PASS 7. 손글씨 공간 원칙(6A)을 코드로 강제한다:
서술형은 1열 전체 폭, 최소 2줄, 이유 설명은 3줄, 표의 서술 셀은 높이를 준다.

실행: ../unified_venv/bin/python worksheets/make_worksheets.py
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).parent
FONT = "맑은 고딕"


def kfont(run):
    """한글이 서양 글꼴로 떨어지지 않게 eastAsia까지 지정한다."""
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return run


# ── 문서 기본 ────────────────────────────────────────────────────────────
def new_doc(title, subtitle, name_line=True):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)          # A4
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.space_after = Pt(4)

    head = doc.add_paragraph()
    run = head.add_run(title)
    run.bold = True
    run.font.size = Pt(17)
    run.font.name = FONT
    kfont(run)

    info = doc.add_paragraph()
    run = info.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    kfont(run)

    if name_line:
        name = doc.add_paragraph()
        run = name.add_run("3학년 　　　반　　　번　　이름 　　　　　　　　　　")
        run.font.size = Pt(11)
        kfont(run)
        bottom_border(name)
        name.paragraph_format.space_after = Pt(10)
    return doc


def bottom_border(par, size=6):
    """밑줄 한 줄. 연속한 문단의 테두리가 같으면 워드가 한 덩어리로 합쳐
    가운데 선을 지워 버리므로, w:between을 함께 넣어 줄 수를 유지한다."""
    from docx.oxml import OxmlElement
    pPr = par._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for tag in ("w:bottom", "w:between"):
        edge = OxmlElement(tag)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "1")
        edge.set(qn("w:color"), "999999")
        borders.append(edge)
    pPr.append(borders)


def shade(cell, hexcolor):
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    tcPr.append(el)


def section(doc, text):
    par = doc.add_paragraph()
    par.paragraph_format.keep_with_next = True
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    kfont(run)
    bottom_border(par, size=12)


def question(doc, text, bold_head=None):
    par = doc.add_paragraph()
    par.paragraph_format.keep_with_next = True
    par.paragraph_format.space_before = Pt(8)
    if bold_head:
        run = par.add_run(bold_head + " ")
        run.bold = True
        kfont(run)
    run = par.add_run(text)
    kfont(run)
    return par


def note(doc, text):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    kfont(run)


def answer_lines(doc, n=2):
    """손글씨 답안 줄. 한 줄에 학생 글씨가 겹치지 않을 높이를 준다."""
    for i in range(n):
        par = doc.add_paragraph()
        par.paragraph_format.keep_with_next = i < n - 1
        par.paragraph_format.space_before = Pt(10)
        par.paragraph_format.space_after = Pt(6)
        bottom_border(par)


def choice(doc, options):
    par = doc.add_paragraph()
    run = par.add_run("　　".join("□ " + o for o in options))
    kfont(run)
    return par


def no_split(row, header=False):
    from docx.oxml import OxmlElement
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))
    if header:
        trPr.append(OxmlElement("w:tblHeader"))


def table(doc, headers, rows, row_cm=1.6, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        kfont(run)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(cell, "EFEFEF")
    no_split(tbl.rows[0], header=True)
    for row in rows:
        cells = tbl.add_row().cells
        tbl.rows[-1].height = Cm(row_cm)
        tbl.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        no_split(tbl.rows[-1])
        for i, text in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            kfont(run)
    # 표는 통째로 한 쪽에 — 마지막 행만 다음 장으로 넘어가면 수업 중 쓰기 어렵다
    for row in tbl.rows[:-1]:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.keep_with_next = True

    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl


def key_item(doc, number, answer, misconception=None, prompt=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    run = par.add_run(number + "　")
    run.bold = True
    kfont(run)
    run = par.add_run(answer)
    kfont(run)
    for label, text in (("자주 나오는 오개념", misconception), ("교사 발문", prompt)):
        if not text:
            continue
        sub = doc.add_paragraph()
        sub.paragraph_format.left_indent = Cm(0.8)
        run = sub.add_run(f"· {label}: ")
        run.bold = True
        run.font.size = Pt(10)
        kfont(run)
        run = sub.add_run(text)
        run.font.size = Pt(10)
        kfont(run)


# ── 15강 학생용 ──────────────────────────────────────────────────────────
def lesson15_student():
    doc = new_doc("15강 학습지 — 전기 에너지의 발생",
                  "[9과22-02] 자석의 운동에 의해 전류가 발생하는 현상을 관찰하고, "
                  "역학적 에너지가 전기 에너지로 전환됨을 설명할 수 있다.")

    question(doc, "건전지가 없는 흔들이 손전등은, 흔들기만 했는데 어떻게 불이 켜질까?",
             "오늘의 질문")

    section(doc, "A. 먼저 생각하기 — 실험 전에 씁니다")
    question(doc, "자석을 코일 안에 넣고 가만히 두면 검류계 바늘이 움직일까요?", "A-1")
    choice(doc, ["움직인다", "움직이지 않는다", "잘 모르겠다"])
    question(doc, "그렇게 생각한 까닭을 쓰세요.")
    answer_lines(doc, 3)
    question(doc, "자석을 더 빠르게 움직이면 무엇이 달라질 것 같나요?", "A-2")
    answer_lines(doc, 2)
    note(doc, "※ 예상이 틀려도 괜찮습니다. 지운 흔적도 그대로 두세요.")

    section(doc, "B. 관찰 기록 — 자석과 코일 애플릿")
    question(doc, "자석을 아래와 같이 움직이며 관찰한 것을 채우세요.")
    table(doc,
          ["자석의 상태", "코일을 통과하는 자기장은 어떻게 변하는가", "검류계 바늘"],
          [["코일 쪽으로 가까이 갈 때", "", ""],
           ["코일 안에 정지해 있을 때", "", ""],
           ["코일 한가운데를 지날 때", "", ""],
           ["코일에서 멀어질 때", "", ""],
           ["같은 구간을 더 빠르게", "", ""]],
          row_cm=1.5, widths=[4.6, 7.4, 5.0])
    question(doc, "자석이 코일 한가운데를 지날 때 바늘이 오히려 작아진 까닭은 무엇일까요?", "B-2")
    answer_lines(doc, 3)

    section(doc, "C. 관찰에서 규칙 찾기")
    question(doc, "전류가 흐르는 조건을 한 문장으로 쓰세요. "
                  "(‘자석이 움직일 때’라고 쓰면 B의 세 번째 줄을 설명할 수 없습니다.)", "C-1")
    answer_lines(doc, 2)
    question(doc, "유도 전류의 방향을 반대로 바꾸는 조작을 두 가지 쓰세요.", "C-2")
    answer_lines(doc, 2)
    question(doc, "유도 전류를 더 세게 만드는 방법을 세 가지 쓰세요.", "C-3")
    answer_lines(doc, 3)

    # 자연스러운 흐름에 맡긴다 — 강제 페이지 나눔은 마지막 쪽을 비운다

    section(doc, "D. 핵심 개념 정리")
    question(doc, "빈칸을 채우세요.", "D-1")
    question(doc, "코일을 통과하는 　　　　　　　이(가) 변하면 코일에 　　　　　　　이(가) 생기고, "
                  "코일이 닫힌 회로를 이루면 　　　　　　　이(가) 흐른다. 이것을 　　　　　　　(이)라 한다.")
    doc.add_paragraph()
    question(doc, "“움직이기만 하면 전류가 흐른다”는 설명은 왜 정확하지 않을까요? "
                  "반례를 하나 들어 설명하세요.", "D-2")
    answer_lines(doc, 3)

    section(doc, "E. 발전기와 에너지")
    question(doc, "코일이 한 바퀴 도는 동안, 유도 전류가 0이 되는 순간은 코일 면이 어떤 상태일 때인가요?", "E-1")
    answer_lines(doc, 2)
    question(doc, "반대로 가장 커지는 순간은 어떤 상태일 때인가요?", "E-2")
    answer_lines(doc, 2)
    question(doc, "자전거 전조등을 켜면 페달이 무거워집니다. 에너지 전환으로 설명하세요.", "E-3")
    answer_lines(doc, 3)
    question(doc, "에너지 전환 사슬을 채우세요.", "E-4")
    question(doc, "　　　　　　　 에너지　→　　　　　　　 에너지　→　빛에너지 · 열에너지")
    doc.add_paragraph()

    section(doc, "F. 생활 적용")
    question(doc, "다음 장치에서 ‘무엇이 변해서’ 전류가 생기는지 한 줄씩 쓰세요.", "F-1")
    table(doc,
          ["장치", "무엇이 변해서 전류가 생기는가"],
          [["다이내믹 마이크", ""], ["인덕션 레인지", ""], ["교통카드", ""]],
          row_cm=1.5, widths=[5.0, 12.0])
    question(doc, "자석과 코일이 같은 속도로 나란히 함께 움직이면 전류가 흐를까요? 까닭도 쓰세요.", "F-2")
    answer_lines(doc, 3)

    section(doc, "Exit Ticket")
    question(doc, "오늘 배운 것 중에서 A에 적었던 처음 예상과 달랐던 점을 한 가지 쓰세요.")
    answer_lines(doc, 2)
    return doc


# ── 16강 학생용 ──────────────────────────────────────────────────────────
def lesson16_student():
    doc = new_doc("16강 학습지 — 전기 에너지의 전환과 이용",
                  "[9과22-03] 가정에서 전기 에너지가 다양한 형태의 에너지로 전환되는 예를 들고, "
                  "이를 소비 전력과 관련지어 설명할 수 있다.")

    question(doc, "소비 전력이 큰 제품은 언제나 전기 에너지를 더 많이 사용할까?", "오늘의 질문")

    section(doc, "A. 먼저 생각하기 — 자료를 보기 전에 씁니다")
    question(doc, "1800 W 드라이어를 10분, 50 W 선풍기를 6시간 사용했습니다. "
                  "전기 에너지를 더 많이 사용한 쪽은?", "A-1")
    choice(doc, ["드라이어", "선풍기", "아직 판단할 수 없다"])
    question(doc, "그렇게 생각한 까닭을 쓰세요.")
    answer_lines(doc, 3)
    question(doc, "선풍기에서 전기 에너지는 운동 에너지로만 전환될까요?", "A-2")
    answer_lines(doc, 2)

    section(doc, "B. 관찰 기록 1 — 가전제품 에너지 전환 탐색기")
    question(doc, "제품을 세 개 골라 표를 채우세요. 증거 칸에는 "
                  "‘밝다 · 뜨겁다 · 움직인다 · 소리가 난다’처럼 확인할 수 있는 것을 씁니다.")
    table(doc,
          ["제품", "주된 목적이 되는 에너지", "함께 나타나는 에너지", "그렇게 판단한 관찰 증거"],
          [["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
          row_cm=1.9, widths=[3.2, 3.6, 3.6, 6.6])
    question(doc, "열이 ‘목적’인 제품과 열이 ‘함께 나타나는’ 제품을 하나씩 쓰고, "
                  "무엇이 그 차이를 정하는지 쓰세요.", "B-2")
    answer_lines(doc, 3)

    section(doc, "C. 관찰 기록 2 — 전력·시간 미터")
    question(doc, "조건을 하나씩만 바꾸며 전력량이 어떻게 되는지 기록하세요.")
    table(doc,
          ["바꾼 조건", "소비 전력 (W)", "사용 시간", "전력량 (Wh)", "어떻게 되었는가"],
          [["소비 전력만 늘림", "", "", "", ""],
           ["사용 시간만 늘림", "", "", "", ""],
           ["전력 ↑ · 시간 ↓", "", "", "", ""]],
          row_cm=1.7, widths=[3.4, 2.6, 2.6, 2.6, 5.8])
    question(doc, "위 결과에서 찾은 규칙을 두 문장으로 쓰세요.", "C-2")
    answer_lines(doc, 3)

    # 자연스러운 흐름에 맡긴다 — 강제 페이지 나눔은 마지막 쪽을 비운다

    section(doc, "D. 핵심 개념 정리")
    question(doc, "소비 전력과 전력량의 차이를 자기 말로 설명하고, 각각의 단위를 쓰세요.", "D-1")
    answer_lines(doc, 3)
    question(doc, "수도꼭지에 비유할 때 W와 Wh는 각각 무엇에 해당하나요?", "D-2")
    answer_lines(doc, 2)

    section(doc, "E. 라벨 읽기와 비교")
    question(doc, "제품 라벨에서 읽은 값을 적고, 같은 시간 사용한다는 조건에서 순서를 매기세요.", "E-1")
    table(doc,
          ["제품", "소비 전력 (W)", "같은 시간 쓸 때 사용량 순서 (1~3)"],
          [["LED 스탠드", "", ""], ["선풍기", "", ""], ["헤어드라이어", "", ""]],
          row_cm=1.4, widths=[5.0, 4.0, 8.0])
    question(doc, "1800 W로 10분, 50 W로 6시간 사용했을 때의 전력량을 각각 구하세요. "
                  "계산 과정도 남기세요.", "E-2")
    answer_lines(doc, 3)
    question(doc, "A-1에 적었던 처음 예상과 비교하면 무엇을 고쳐야 하나요?", "E-3")
    answer_lines(doc, 2)
    question(doc, "전력량이 200 Wh가 되는 ‘소비 전력 × 사용 시간’ 조합을 두 가지 만드세요.", "E-4")
    answer_lines(doc, 2)

    section(doc, "F. 한 단계 더 생각하기")
    question(doc, "“10 W 스탠드가 1800 W 드라이어보다 효율이 좋다”고 바로 말할 수 있나요? "
                  "까닭을 쓰세요.", "F-1")
    answer_lines(doc, 3)
    question(doc, "“W가 가장 큰 제품을 쓰지 말자”보다 더 타당한 절약 원칙을 하나 쓰고, "
                  "전력과 사용 시간 중 무엇을 근거로 삼았는지 밝히세요.", "F-2")
    answer_lines(doc, 3)

    section(doc, "Exit Ticket")
    question(doc, "오늘 집에서 가전제품 라벨을 본다면 어떤 두 정보를 함께 확인하겠습니까? "
                  "그 이유를 한 문장으로 쓰세요.")
    answer_lines(doc, 2)
    return doc


# ── 교사용 정답·해설 ─────────────────────────────────────────────────────
def lesson15_teacher():
    doc = new_doc("15강 학습지 교사용 — 전기 에너지의 발생",
                  "학생 배부용이 아닙니다. 정답 예시 · 오개념 · 발문", name_line=False)

    section(doc, "A. 먼저 생각하기")
    key_item(doc, "A-1", "움직이지 않는다. 코일을 통과하는 자기장이 변하지 않기 때문이다. "
             "다만 이 단계에서는 정답을 확인해 주지 않는다.",
             "‘자석이 코일 안에 있으니 강한 자기장이 있어서 전류가 흐른다’ — 세기와 변화를 혼동한다.",
             "지금 답이 맞는지가 아니라, 무엇을 보고 그렇게 생각했는지 말하게 한다.")
    key_item(doc, "A-2", "더 세게(크게) 흔들린다 / 바늘이 더 많이 움직인다 수준이면 충분하다.")

    section(doc, "B. 관찰 기록")
    key_item(doc, "표", "가까이 갈 때: 점점 세진다 · 한쪽으로 크게 / 안에 정지: 일정하다 · 0 / "
             "한가운데 통과: 거의 그대로(변화 작음) · 작아짐 / 멀어질 때: 점점 약해진다 · 반대쪽으로 / "
             "빠르게: 같은 변화를 짧은 시간에 · 더 크게",
             "‘한가운데에서 자기장이 가장 세니까 전류도 최대’ — 세기와 변화율의 혼동.")
    key_item(doc, "B-2", "코일 한가운데에서는 통과하는 자기장이 최대이면서 거의 변하지 않기 때문이다. "
             "전류는 자기장의 값이 아니라 그 변화에 따라 생긴다.",
             None,
             "‘가장 센 곳인데 왜 작아졌을까?’ 이 한 문장이 이번 차시의 분기점이다.")

    section(doc, "C. 규칙 찾기")
    key_item(doc, "C-1", "코일을 통과하는 자기장이 변할 때 유도 전압이 생기고, 닫힌 회로이면 유도 전류가 흐른다.",
             "‘자석이 움직이면’이라고 쓰면 F-2(나란히 이동)를 설명하지 못한다.")
    key_item(doc, "C-2", "① 가까이 갈 때와 멀어질 때를 바꾼다 ② 자석의 극을 뒤집는다. "
             "둘 다 바꾸면 방향은 원래대로 돌아온다.")
    key_item(doc, "C-3", "더 빠르게 움직인다 / 더 센 자석을 쓴다 / 코일을 더 많이 감는다. "
             "(회로의 저항 등 다른 조건은 같다고 본 단순 비교다.)")

    section(doc, "D. 핵심 개념")
    key_item(doc, "D-1", "자기장 · 유도 전압 · 유도 전류 · 전자기 유도")
    key_item(doc, "D-2", "자석과 코일이 나란히 함께 움직이면 거리가 그대로여서 통과 자기장이 변하지 않고, "
             "따라서 전류도 흐르지 않는다. 조건은 움직임이 아니라 통과 자기장의 변화다.")

    section(doc, "E. 발전기와 에너지")
    key_item(doc, "E-1", "코일 면이 자기장에 수직일 때. 통과하는 자기장이 최대라서 변화가 가장 작다.",
             "‘코일이 멈춘 순간’이라고 답하기 쉽다. 코일은 계속 돌고 있다.")
    key_item(doc, "E-2", "코일 면이 자기장과 나란할 때. 통과 자기장이 가장 빠르게 바뀐다.")
    key_item(doc, "E-3", "전기를 만든 만큼 누군가 일을 해야 한다. 전조등을 켜면 회로에 전류가 흐르고, "
             "그 전기 에너지의 근원이 페달을 밟는 역학적 에너지이므로 페달이 무거워진다.",
             "‘공짜로 전기가 생긴다’ 또는 ‘자석이 에너지를 만든다’.")
    key_item(doc, "E-4", "역학적(운동) → 전기 → 빛·열")

    section(doc, "F. 생활 적용 · Exit")
    key_item(doc, "F-1", "마이크: 소리가 진동판과 코일을 흔들어 통과 자기장이 변한다 / "
             "인덕션: 코일의 변하는 자기장이 냄비에 전류를 유도한다 / "
             "교통카드: 단말기의 변하는 자기장이 카드 코일에 전류를 유도한다.")
    key_item(doc, "F-2", "흐르지 않는다. 서로의 거리가 그대로여서 통과 자기장이 변하지 않는다.")
    key_item(doc, "Exit", "A-1을 ‘움직인다’로 적었던 학생이 ‘변화가 있어야 한다’로 고쳐 쓰면 목표 도달로 본다.")
    return doc


def lesson16_teacher():
    doc = new_doc("16강 학습지 교사용 — 전기 에너지의 전환과 이용",
                  "학생 배부용이 아닙니다. 정답 예시 · 오개념 · 발문", name_line=False)

    section(doc, "A. 먼저 생각하기")
    key_item(doc, "A-1", "이 단계에서는 정답을 공개하지 않는다. 세 선택지 모두 받아 준다. "
             "실제로는 둘 다 300 Wh로 같다(E-2에서 확인).",
             "‘W가 크니까 드라이어’ — 전력과 전력량의 혼동. 이번 차시의 핵심 표적이다.",
             "무엇을 보고 골랐는지, 판단하려면 어떤 정보가 하나 더 필요한지 묻는다.")
    key_item(doc, "A-2", "아니다. 날개와 공기의 운동 외에 소리와 모터의 열도 함께 나타난다.")

    section(doc, "B. 에너지 전환")
    key_item(doc, "표", "전등: 빛 / 열 · 선풍기: 운동 / 소리·열 · 드라이어: 열·바람의 운동 / 소리 · "
             "스피커: 소리 / 진동·열 · 밥솥: 열 / 소리 · 세탁기: 운동 / 소리·열",
             "‘가전제품 하나는 한 종류의 에너지만 만든다’.")
    key_item(doc, "B-2", "밥솥·난로·드라이어에서는 열이 목적이고, 선풍기·스피커에서는 함께 나타나는 에너지다. "
             "차이를 정하는 것은 그 제품을 쓰는 목적이다.",
             "‘열은 언제나 낭비다’.")

    section(doc, "C. 전력·시간")
    key_item(doc, "C-2", "① 같은 시간에는 소비 전력이 클수록 전력량이 커진다. "
             "② 같은 소비 전력에서는 사용 시간이 길수록 전력량이 커진다. "
             "(전력과 시간이 모두 다르면 둘을 함께 봐야 한다.)")

    section(doc, "D. 핵심 개념")
    key_item(doc, "D-1", "소비 전력은 전기 에너지를 사용하는 빠르기이고 단위는 W다. "
             "전력량은 일정 시간 동안 사용한 전기 에너지의 양이고 단위는 Wh 또는 kWh다.",
             "‘kWh는 kW/h다’ — 곱과 나눗셈의 혼동. 면적 모형으로 되돌아간다.")
    key_item(doc, "D-2", "W는 물이 나오는 빠르기, Wh는 그동안 통에 모인 물의 전체 양.",
             None, "비유의 한계도 함께 말한다. 전기가 관 속의 물처럼 저장되어 흐르는 것은 아니다.")

    section(doc, "E. 라벨과 비교")
    key_item(doc, "E-1", "LED 스탠드 10 W(3위) · 선풍기 50 W(2위) · 드라이어 1800 W(1위). "
             "‘같은 시간’이라는 조건이 있을 때만 이렇게 줄 세울 수 있다.")
    key_item(doc, "E-2", "드라이어: 1800 W × (10/60)시간 = 300 Wh, 선풍기: 50 W × 6시간 = 300 Wh. 둘이 같다.",
             None, "계산보다 먼저 애플릿의 직사각형 넓이를 눈으로 비교하게 한다.")
    key_item(doc, "E-3", "‘소비 전력이 크면 전기를 더 많이 쓴다’를 "
             "‘같은 시간이라면’이라는 조건이 붙어야 한다로 고치면 도달로 본다.")
    key_item(doc, "E-4", "예: 100 W × 2시간, 200 W × 1시간, 400 W × 30분, 1200 W × 10분.")

    section(doc, "F. 한 단계 더 · Exit")
    key_item(doc, "F-1", "말할 수 없다. 두 제품은 목적과 유효 출력이 달라 비교 기준이 없다. "
             "효율은 같은 기능·비슷한 용량·같은 시험 조건에서만 비교한다.",
             "‘소비 전력이 작으면 효율이 좋다’.")
    key_item(doc, "F-2", "필요 없을 때 사용 시간을 줄인다 / 같은 품목·용량끼리 더 효율적인 제품을 고른다 / "
             "대기 전력을 확인한다. 근거가 전력인지 시간인지 밝혔는지를 본다.")
    key_item(doc, "Exit", "소비 전력(W)과 사용 시간(또는 연간 소비 전력량 kWh) 두 가지를 함께 확인하겠다는 답이면 도달.")
    return doc


if __name__ == "__main__":
    for name, build in [
        ("15강_학습지_학생용", lesson15_student),
        ("15강_학습지_교사용", lesson15_teacher),
        ("16강_학습지_학생용", lesson16_student),
        ("16강_학습지_교사용", lesson16_teacher),
    ]:
        path = OUT / f"{name}.docx"
        build().save(path)
        print("saved", path.name)
